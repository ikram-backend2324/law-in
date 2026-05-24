from django.shortcuts import render, redirect
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib import messages
from django.http import HttpResponse
from .parsers import parse_file
import openpyxl
import csv
import io


@staff_member_required
def bulk_upload_view(request):
    if request.method == 'POST':
        uploaded = request.FILES.get('bulk_file')
        if not uploaded:
            messages.error(request, 'Please select a file to upload.')
            return redirect('bulk_upload')
        try:
            tests_count, q_count, errors = parse_file(uploaded, uploaded.name, request.user)
            if errors:
                for e in errors[:10]:
                    messages.warning(request, e)
            if tests_count > 0 or q_count > 0:
                messages.success(request, f'✅ Successfully imported {tests_count} test(s) and {q_count} question(s).')
            else:
                messages.warning(request, 'No data was imported. Please check your file format.')
        except Exception as e:
            messages.error(request, f'Error parsing file: {str(e)}')
        return redirect('bulk_upload')

    return render(request, 'tests_app/bulk_upload.html')


@staff_member_required
def download_template(request, fmt):
    headers = ['test_title', 'question_text', 'option_a', 'option_b', 'option_c', 'option_d', 'correct_answer']
    sample = [
        ['Constitutional Law', 'What is the supreme law of the land?', 'The President', 'The Constitution', 'Congress', 'The Courts', 'B'],
        ['Constitutional Law', 'How many branches of government are there?', 'Two', 'Three', 'Four', 'Five', 'B'],
        ['Civil Law', 'What does "plaintiff" mean?', 'The judge', 'The defendant', 'The person who files a lawsuit', 'The lawyer', 'C'],
    ]

    if fmt == 'xlsx':
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Template'
        ws.append(headers)
        for row in sample:
            ws.append(row)
        # Style header
        from openpyxl.styles import Font, PatternFill
        for cell in ws[1]:
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor='1a3a5c')
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        resp['Content-Disposition'] = 'attachment; filename="law_in_template.xlsx"'
        return resp

    elif fmt == 'csv':
        resp = HttpResponse(content_type='text/csv')
        resp['Content-Disposition'] = 'attachment; filename="law_in_template.csv"'
        writer = csv.writer(resp)
        writer.writerow(headers)
        for row in sample:
            writer.writerow(row)
        return resp

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading('Law In — Bulk Upload Template', 0)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
        for row in sample:
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = val
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp['Content-Disposition'] = 'attachment; filename="law_in_template.docx"'
        return resp

    elif fmt == 'pdf':
        # Return a simple text file with instructions since PDF generation requires reportlab
        content = "Law In Bulk Upload Template\n"
        content += "=" * 60 + "\n\n"
        content += "Required columns (in order):\n"
        for h in headers:
            content += f"  - {h}\n"
        content += "\nSample rows:\n"
        for row in sample:
            content += "  " + " | ".join(row) + "\n"
        content += "\nNote: For PDF uploads, your PDF must contain a table with these exact column headers.\n"
        resp = HttpResponse(content, content_type='text/plain')
        resp['Content-Disposition'] = 'attachment; filename="law_in_pdf_instructions.txt"'
        return resp

    return HttpResponse('Invalid format', status=400)

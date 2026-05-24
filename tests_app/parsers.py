import csv
import io
import re
import openpyxl
import pdfplumber
import docx as python_docx
from .models import Test, Question, Answer


# ─────────────────────────────────────────────
#  Shared DB writer
# ─────────────────────────────────────────────

def _create_from_rows(rows, created_by):
    tests_created = {}
    questions_created = 0
    errors = []

    for i, row in enumerate(rows, start=2):
        try:
            test_title    = str(row.get('test_title', '')).strip()
            question_text = str(row.get('question_text', '')).strip()
            option_a      = str(row.get('option_a', '')).strip()
            option_b      = str(row.get('option_b', '')).strip()
            option_c      = str(row.get('option_c', '')).strip()
            option_d      = str(row.get('option_d', '')).strip()
            correct       = str(row.get('correct_answer', '')).strip().upper()

            if not test_title or not question_text:
                continue
            if correct not in ('A', 'B', 'C', 'D'):
                errors.append(f"Row {i}: correct_answer must be A, B, C or D (got '{correct}')")
                continue

            if test_title not in tests_created:
                test, _ = Test.objects.get_or_create(
                    title=test_title,
                    defaults={'created_by': created_by, 'is_active': True}
                )
                tests_created[test_title] = test
            else:
                test = tests_created[test_title]

            q_order  = test.questions.count() + 1
            question = Question.objects.create(test=test, text=question_text, order=q_order)
            questions_created += 1

            for label, text in [('A', option_a), ('B', option_b), ('C', option_c), ('D', option_d)]:
                if text:
                    Answer.objects.create(question=question, label=label, text=text, is_correct=(label == correct))
        except Exception as e:
            errors.append(f"Row {i}: {str(e)}")

    return len(tests_created), questions_created, errors


# ─────────────────────────────────────────────
#  Free-text parser  (docx / pdf plain text)
#
#  Supported format (flexible, auto-detected):
#
#  Test: My Test Title
#  1. What is X?
#  A) Option one
#  B) Option two
#  C) Option three
#  D) Option four
#  Answer: B
#
# ─────────────────────────────────────────────

def _parse_freetext(text, created_by):
    rows = []
    current_test  = 'Uploaded Test'
    current_q     = None
    current_opts  = {}
    current_ans   = None

    def flush():
        nonlocal current_q, current_opts, current_ans
        if current_q and len(current_opts) >= 2 and current_ans:
            rows.append({
                'test_title':    current_test,
                'question_text': current_q,
                'option_a':      current_opts.get('A', ''),
                'option_b':      current_opts.get('B', ''),
                'option_c':      current_opts.get('C', ''),
                'option_d':      current_opts.get('D', ''),
                'correct_answer': current_ans,
            })
        current_q   = None
        current_opts = {}
        current_ans  = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Test title line:  "Test: Foo"  or  "Test Title: Foo"
        m = re.match(r'^[Tt]est(?:\s+[Tt]itle)?[:\-]\s*(.+)', line)
        if m:
            flush()
            current_test = m.group(1).strip()
            continue

        # Answer line:  "Answer: B"  or  "Correct: B"  or  "Javob: B"
        m = re.match(r'^(?:Answer|Correct|Javob|answer|correct)[:\s]+([A-Da-d])', line)
        if m:
            current_ans = m.group(1).upper()
            continue

        # Option line:  "A) ..."  "A. ..."  "A: ..."  "(A) ..."
        m = re.match(r'^\(?([A-Da-d])[).\s:]\s*(.+)', line)
        if m:
            current_opts[m.group(1).upper()] = m.group(2).strip()
            continue

        # Question line: starts with a number  "1. ..."  "1) ..."  "Q1: ..."
        m = re.match(r'^(?:[Qq]?\d+[.):\s]+)\s*(.+)', line)
        if m:
            flush()
            current_q = m.group(1).strip()
            continue

        # If we already have a question but this line doesn't match any pattern,
        # append it to the question text (multi-line questions)
        if current_q and not current_opts:
            current_q += ' ' + line

    flush()
    return _create_from_rows(rows, created_by)


# ─────────────────────────────────────────────
#  Format parsers
# ─────────────────────────────────────────────

def parse_xlsx(file_obj, created_by):
    wb = openpyxl.load_workbook(file_obj)
    ws = wb.active
    rows_data, headers = [], None
    for row in ws.iter_rows(values_only=True):
        if headers is None:
            headers = [str(c).strip().lower() if c else '' for c in row]
            continue
        if all(c is None for c in row):
            continue
        rows_data.append(dict(zip(headers, [str(c).strip() if c is not None else '' for c in row])))
    return _create_from_rows(rows_data, created_by)


def parse_csv(file_obj, created_by):
    text = file_obj.read().decode('utf-8-sig')
    reader = csv.DictReader(io.StringIO(text))
    rows_data = [{k.strip().lower(): v.strip() for k, v in row.items()} for row in reader]
    return _create_from_rows(rows_data, created_by)


def parse_pdf(file_obj, created_by):
    """Try table extraction first; fall back to free-text parsing."""
    rows_data = []
    full_text = []

    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if not table:
                    continue
                headers = None
                for row in table:
                    if headers is None:
                        headers = [str(c).strip().lower() if c else '' for c in row]
                        continue
                    if not any(row):
                        continue
                    rows_data.append(dict(zip(headers, [str(c).strip() if c else '' for c in row])))
            # Collect plain text for fallback
            full_text.append(page.extract_text() or '')

    if rows_data:
        return _create_from_rows(rows_data, created_by)
    # Fallback: parse as free text
    return _parse_freetext('\n'.join(full_text), created_by)


def parse_docx(file_obj, created_by):
    """Try table extraction first; fall back to free-text parsing."""
    doc = python_docx.Document(file_obj)

    # Try tables
    rows_data = []
    for table in doc.tables:
        headers = None
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if headers is None:
                headers = [c.lower() for c in cells]
                continue
            if not any(cells):
                continue
            rows_data.append(dict(zip(headers, cells)))

    if rows_data:
        return _create_from_rows(rows_data, created_by)

    # Fallback: join all paragraph text and parse free-text
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    return _parse_freetext(full_text, created_by)


def parse_file(file_obj, filename, created_by):
    ext = filename.lower().rsplit('.', 1)[-1]
    if ext == 'xlsx':
        return parse_xlsx(file_obj, created_by)
    elif ext == 'csv':
        return parse_csv(file_obj, created_by)
    elif ext == 'pdf':
        return parse_pdf(file_obj, created_by)
    elif ext == 'docx':
        return parse_docx(file_obj, created_by)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")